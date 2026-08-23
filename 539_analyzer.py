import os
import csv
import sys
import datetime
from collections import Counter

def load_data():
    """
    載入 2025 與 2026 今彩539開獎數據並進行整合排序
    """
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    csv_files = [
        os.path.join(base_dir, "2025", "今彩539_2025.csv"),
        os.path.join(base_dir, "2026", "今彩539_2026.csv")
    ]
    
    all_draws = []
    
    for file_path in csv_files:
        if not os.path.exists(file_path):
            print(f"警告: 找不到檔案 {file_path}")
            continue
            
        # 嘗試以不同的編碼讀取
        encodings = ['utf-8-sig', 'utf-8', 'big5', 'gbk']
        content_lines = None
        for enc in encodings:
            try:
                with open(file_path, mode='r', encoding=enc) as f:
                    content_lines = list(csv.reader(f))
                break
            except (UnicodeDecodeError, LookupError):
                continue
                
        if content_lines is None:
            print(f"錯誤: 無法解析檔案 {file_path} 的編碼。")
            continue
            
        for row in content_lines[1:]:
            if not row or len(row) < 3:
                continue
            date_str = row[2].strip()
            if not date_str or ('/' not in date_str and '-' not in date_str):
                continue
                
            non_empty_nums = [x.strip() for x in row[3:] if x.strip()]
            if len(non_empty_nums) < 5:
                continue
            try:
                nums = [int(x) for x in non_empty_nums[-5:]]
            except ValueError:
                continue
                
            period = None
            if row[1].strip():
                try:
                    period = int(row[1].strip())
                except ValueError:
                    pass
                    
            nums.sort()
            all_draws.append({
                "period": period,
                "date": date_str,
                "nums": nums
            })
            
    def get_date_key(draw):
        date_str = draw["date"]
        sep = '/' if '/' in date_str else '-'
        parts = date_str.split(sep)
        return (int(parts[0]), int(parts[1]), int(parts[2]))
        
    all_draws.sort(key=get_date_key)
    
    for idx, draw in enumerate(all_draws):
        if draw["period"] is None:
            if idx > 0 and all_draws[idx-1]["period"] is not None:
                draw["period"] = all_draws[idx-1]["period"] + 1
            else:
                draw["period"] = 114000000 + idx
                
    return all_draws


def run_drag_analysis(draws, selected_nums, threshold=2):
    """
    拖牌分析：尋找歷史上包含 >= threshold 個 selected_nums 的期數，統計其下一期的開獎號碼頻率 (1-39)
    """
    matches = []
    next_period_nums = []
    
    selected_set = set(selected_nums)
    
    for i in range(len(draws) - 1):
        curr_draw = draws[i]
        next_draw = draws[i+1]
        
        intersect = selected_set.intersection(curr_draw["nums"])
        if len(intersect) >= threshold:
            matches.append({
                "curr": curr_draw,
                "next": next_draw,
                "intersect": sorted(list(intersect))
            })
            next_period_nums.extend(next_draw["nums"])
            
    freq = Counter(next_period_nums)
    
    for n in range(1, 40):
        if n not in freq:
            freq[n] = 0
            
    return matches, freq

def run_hot_cold_analysis(draws, window=30):
    """
    冷熱門分析：計算最近 window 期內各號碼出現頻率 (1-39)
    """
    recent_draws = draws[-window:] if len(draws) >= window else draws
    all_nums = []
    for d in recent_draws:
        all_nums.extend(d["nums"])
        
    counts = Counter(all_nums)
    
    for n in range(1, 40):
        if n not in counts:
            counts[n] = 0
            
    sorted_counts = sorted(counts.items(), key=lambda x: x[1], reverse=True)
    
    # 539 總共 39 個號碼，取前 8 熱門，後 8 冷門
    hot_nums = [n for n, c in sorted_counts[:8]]
    cold_nums = [n for n, c in sorted_counts[-8:]]
    
    return counts, hot_nums, cold_nums

def run_gap_analysis(draws):
    """
    未開期數間隔分析 (1-39)
    """
    gap_stats = {n: {"current_gap": 0, "gaps": [], "last_idx": -1} for n in range(1, 40)}
    
    for idx, d in enumerate(draws):
        for n in d["nums"]:
            if gap_stats[n]["last_idx"] != -1:
                gap = idx - gap_stats[n]["last_idx"] - 1
                gap_stats[n]["gaps"].append(gap)
            gap_stats[n]["last_idx"] = idx
            
    total_periods = len(draws)
    
    for n in range(1, 40):
        last_idx = gap_stats[n]["last_idx"]
        if last_idx == -1:
            gap_stats[n]["current_gap"] = total_periods
        else:
            gap_stats[n]["current_gap"] = total_periods - 1 - last_idx
            
        gaps = gap_stats[n]["gaps"]
        if gaps:
            gap_stats[n]["avg_gap"] = round(sum(gaps) / len(gaps), 1)
            gap_stats[n]["max_gap"] = max(gaps)
        else:
            gap_stats[n]["avg_gap"] = total_periods
            gap_stats[n]["max_gap"] = total_periods
            
    return gap_stats

def backtest_strategies(draws, test_window=20, threshold=2):
    """
    回測系統：預測 5 個號碼
    1. 策略 A：純拖牌頻率 Top 5
    2. 策略 B：綜合指標 (拖牌權重 + 遺漏值權重)
    """
    if len(draws) < test_window + 50:
        return None
        
    results_A = []
    results_B = []
    
    start_idx = len(draws) - test_window
    
    for i in range(start_idx, len(draws)):
        hist_draws = draws[:i]
        actual_draw = draws[i]
        
        prev_draw = hist_draws[-1]
        selected = prev_draw["nums"]
        
        _, drag_freq = run_drag_analysis(hist_draws, selected, threshold)
        gap_stats = run_gap_analysis(hist_draws)
        
        top_A = sorted(range(1, 40), key=lambda x: (drag_freq[x], -x), reverse=True)[:5]
        
        scores = {}
        max_drag = max(drag_freq.values()) if drag_freq.values() else 1
        
        for n in range(1, 40):
            s_drag = drag_freq[n] / max_drag
            current_gap = gap_stats[n]["current_gap"]
            avg_gap = gap_stats[n]["avg_gap"] if gap_stats[n]["avg_gap"] > 0 else 7.8
            s_gap = min(current_gap / avg_gap, 2.0) / 2.0
            
            scores[n] = 0.6 * s_drag + 0.4 * s_gap
            
        top_B = sorted(range(1, 40), key=lambda x: scores[x], reverse=True)[:5]
        
        hits_A = len(set(top_A).intersection(actual_draw["nums"]))
        hits_B = len(set(top_B).intersection(actual_draw["nums"]))
        
        results_A.append({"period": actual_draw["period"], "hits": hits_A})
        results_B.append({"period": actual_draw["period"], "hits": hits_B})
        
    return {
        "strategy_A": results_A,
        "strategy_B": results_B
    }

def generate_docx(latest_draw, selected_nums, threshold, top_recommended, scores, drag_freq, gap_stats, hot_nums, cold_nums, matches, hot_cold_counts, backtest_data, file_prefix):
    """
    將分析結果匯出為 Word .docx 檔案
    """
    try:
        import docx
        from docx.shared import Pt, RGBColor
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
    except ImportError:
        print("未安裝 python-docx，無法生成 .docx 報告。")
        return

    doc = docx.Document()
    
    title = doc.add_heading('今彩539數據分析與預測報告', level=1)
    title.alignment = 1
    
    p_meta = doc.add_paragraph()
    p_meta.add_run(f"報告生成時間: {latest_draw['date']} (資料庫最新期別: {latest_draw['period']})\n").italic = True
    p_meta.add_run("本期分析輸入 (最新獎號): " + " ".join(f"{n:02d}" for n in selected_nums) + "\n").italic = True
    p_meta.add_run(f"匹配拖牌門檻: >= {threshold} 星\n").italic = True

    doc.add_heading('1. 預測推薦號碼 (綜合評分 Top 10)', level=2)
    doc.add_paragraph("結合了拖牌出現頻率 (權重 60%) 與 遺漏值超限程度 (權重 40%) 的綜合評估推薦：")
    
    table_rec = doc.add_table(rows=1, cols=7)
    table_rec.style = 'Table Grid'
    headers_rec = ["推薦排名", "號碼", "綜合分數", "拖牌次數", "目前遺漏期數", "平均間隔期數", "狀態標籤"]
    hdr_cells = table_rec.rows[0].cells
    for idx, name in enumerate(headers_rec):
        hdr_cells[idx].text = name
        if hdr_cells[idx].paragraphs[0].runs:
            hdr_cells[idx].paragraphs[0].runs[0].font.bold = True
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F2F2"/>')
        hdr_cells[idx]._tc.get_or_add_tcPr().append(shading_elm)

    for rank, n in enumerate(top_recommended, 1):
        row_cells = table_rec.add_row().cells
        tag = []
        if n in hot_nums:
            tag.append("近期熱門")
        elif n in cold_nums:
            tag.append("近期冷門")
        if gap_stats[n]["current_gap"] > gap_stats[n]["avg_gap"] * 1.5:
            tag.append("遺漏值偏高")
        elif gap_stats[n]["current_gap"] == 0:
            tag.append("上期剛開")
        tag_str = "/".join(tag) if tag else "正常"
        
        row_cells[0].text = str(rank)
        row_cells[1].text = f"{n:02d}"
        if row_cells[1].paragraphs[0].runs:
            row_cells[1].paragraphs[0].runs[0].font.bold = True
        row_cells[2].text = f"{scores[n]:.2f}"
        row_cells[3].text = f"{drag_freq[n]} 次"
        row_cells[4].text = f"{gap_stats[n]['current_gap']} 期"
        row_cells[5].text = f"{gap_stats[n]['avg_gap']} 期"
        row_cells[6].text = tag_str

    # 新增純拖牌前 5 名 (含第 5 名同次數號碼)
    doc.add_paragraph()
    p_pure_hdr = doc.add_paragraph()
    p_pure_hdr.add_run("純拖牌分析前 5 名 (含第 5 名同次數號碼)：").bold = True
    
    sorted_drag = sorted(drag_freq.items(), key=lambda x: x[1], reverse=True)
    cutoff_val = sorted_drag[4][1] if len(sorted_drag) > 4 else 1
    pure_drag_top = [item for item in sorted_drag if item[1] >= cutoff_val and item[1] > 0]
    pure_drag_top.sort(key=lambda x: (x[1], -x[0]), reverse=True)
    
    current_rank = 1
    for idx, (num, count) in enumerate(pure_drag_top):
        if idx > 0 and count < pure_drag_top[idx - 1][1]:
            current_rank = idx + 1
        p_item = doc.add_paragraph(style='List Bullet')
        p_item.add_run(f"第 {current_rank} 名：號碼 ").bold = True
        p_item.add_run(f"{num:02d}").bold = True
        p_item.add_run(f" (歷史拖牌出現 {count} 次)")

    doc.add_heading('2. 歷史拖牌匹配明細 (近 10 次匹配)', level=2)
    doc.add_paragraph(f"尋找歷史上開出與本期相同的獎號中，包含 >= {threshold} 個號碼的期數，其下一期的開獎結果：")
    
    table_match = doc.add_table(rows=1, cols=4)
    table_match.style = 'Table Grid'
    headers_match = ["歷史期別", "開獎日期", "當期開獎號碼", "其下一期開獎號碼"]
    hdr_cells_match = table_match.rows[0].cells
    for idx, name in enumerate(headers_match):
        hdr_cells_match[idx].text = name
        if hdr_cells_match[idx].paragraphs[0].runs:
            hdr_cells_match[idx].paragraphs[0].runs[0].font.bold = True
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F2F2"/>')
        hdr_cells_match[idx]._tc.get_or_add_tcPr().append(shading_elm)

    for m in reversed(matches[-10:]):
        row_cells = table_match.add_row().cells
        row_cells[0].text = str(m["curr"]["period"])
        row_cells[1].text = m["curr"]["date"]
        
        p_curr = row_cells[2].paragraphs[0]
        for n in m["curr"]["nums"]:
            run = p_curr.add_run(f"{n:02d} ")
            if n in selected_nums:
                run.bold = True
                run.font.color.rgb = RGBColor(220, 50, 50)
                
        row_cells[3].text = " ".join(f"{n:02d}" for n in m["next"]["nums"])
        if row_cells[3].paragraphs[0].runs:
            row_cells[3].paragraphs[0].runs[0].font.bold = True

    doc.add_paragraph(f"\n歷史匹配總次數: {len(matches)} 次").italic = True

    doc.add_heading('3. 冷熱門與遺漏值分佈 (1-39 完整表)', level=2)
    p_hc = doc.add_paragraph()
    p_hc.add_run("熱門號碼 (近30期開出次數最高前8名): ").bold = True
    p_hc.add_run(", ".join(f"{n:02d}" for n in hot_nums) + "\n")
    p_hc.add_run("冷門號碼 (近30期開出次數最低後8名): ").bold = True
    p_hc.add_run(", ".join(f"{n:02d}" for n in cold_nums))
    
    table_full = doc.add_table(rows=1, cols=6)
    table_full.style = 'Table Grid'
    headers_full = ["號碼", "拖牌次數", "近30期次數", "目前遺漏期數", "歷史平均間隔", "歷史最大間隔"]
    hdr_cells_full = table_full.rows[0].cells
    for idx, name in enumerate(headers_full):
        hdr_cells_full[idx].text = name
        if hdr_cells_full[idx].paragraphs[0].runs:
            hdr_cells_full[idx].paragraphs[0].runs[0].font.bold = True
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F2F2"/>')
        hdr_cells_full[idx]._tc.get_or_add_tcPr().append(shading_elm)

    for n in range(1, 40):
        row_cells = table_full.add_row().cells
        row_cells[0].text = f"{n:02d}"
        if row_cells[0].paragraphs[0].runs:
            row_cells[0].paragraphs[0].runs[0].font.bold = True
        row_cells[1].text = str(drag_freq[n])
        row_cells[2].text = str(hot_cold_counts[n])
        row_cells[3].text = str(gap_stats[n]["current_gap"])
        row_cells[4].text = str(gap_stats[n]["avg_gap"])
        row_cells[5].text = str(gap_stats[n]["max_gap"])

    doc.add_heading('4. 預測策略歷史回測分析 (近 30 期表現)', level=2)
    doc.add_paragraph("模擬在過去 30 期中，如果每一期我們都使用當時的歷史資料進行預測，並購買推薦的前 5 名號碼，實際的中獎情況：")
    
    if backtest_data:
        hits_A = [r["hits"] for r in backtest_data["strategy_A"]]
        hits_B = [r["hits"] for r in backtest_data["strategy_B"]]
        
        def get_hit_stats_dict(hits):
            total = len(hits)
            h0 = hits.count(0)
            h1 = hits.count(1)
            h2 = hits.count(2)
            h3 = hits.count(3)
            h4_plus = sum(1 for h in hits if h >= 4)
            return {
                "avg": round(sum(hits)/total, 2),
                "rates": [
                    f"中 0 星: {h0/total*100:.1f}% ({h0}次)",
                    f"中 1 星: {h1/total*100:.1f}% ({h1}次)",
                    f"中 2 星: {h2/total*100:.1f}% ({h2}次)",
                    f"中 3 星(肆獎): {h3/total*100:.1f}% ({h3}次)",
                    f"中 4 星以上: {h4_plus/total*100:.1f}% ({h4_plus}次)"
                ]
            }
            
        stats_A = get_hit_stats_dict(hits_A)
        stats_B = get_hit_stats_dict(hits_B)
        
        p_strA = doc.add_paragraph()
        p_strA.add_run("策略 A (純拖牌頻率 Top 5):\n").bold = True
        p_strA.add_run(f"  * 歷史平均每期命中: {stats_A['avg']} 個號碼\n")
        for rate in stats_A["rates"]:
            p_strA.add_run(f"  * {rate}\n")
            
        p_strB = doc.add_paragraph()
        p_strB.add_run("策略 B (綜合評分 Top 5 - 拖牌 60% + 遺漏值 40%):\n").bold = True
        p_strB.add_run(f"  * 歷史平均每期命中: {stats_B['avg']} 個號碼\n")
        for rate in stats_B["rates"]:
            p_strB.add_run(f"  * {rate}\n")
            
        doc.add_paragraph("\n*註：命中次數純以 5 個主區獎號為準。今彩 539 中 2 個號碼即可得 50 元，中 3 個號碼得 300 元。").italic = True
    else:
        doc.add_paragraph("數據庫內期數不足，無法進行 30 期回測。")

    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    docx_path = os.path.join(base_dir, f"{file_prefix}539分析.docx")
    try:
        doc.save(docx_path)
        print(f"成功生成 Word 報告：{docx_path}")
    except PermissionError:
        print(f"錯誤: 無法寫入 Word 報告 {docx_path}。請確認您是否已在 Word 中開啟了該檔案？請將檔案關閉後再重新執行腳本。")

def generate_json(latest_draw, selected_nums, threshold, top_recommended, scores, drag_freq, gap_stats, hot_nums, cold_nums, matches, hot_cold_counts, backtest_data):
    """
    將分析結果產出為 JSON 格式供 Web / PWA 使用
    """
    import json
    
    # 1. 綜合推薦 Top 10
    recommendations = []
    for rank, n in enumerate(top_recommended, 1):
        tag = []
        if n in hot_nums:
            tag.append("近期熱門")
        elif n in cold_nums:
            tag.append("近期冷門")
        if gap_stats[n]["current_gap"] > gap_stats[n]["avg_gap"] * 1.5:
            tag.append("遺漏值偏高")
        elif gap_stats[n]["current_gap"] == 0:
            tag.append("上期剛開")
        
        recommendations.append({
            "rank": rank,
            "num": f"{n:02d}",
            "score": round(scores[n], 2),
            "drag_count": drag_freq[n],
            "current_gap": gap_stats[n]["current_gap"],
            "avg_gap": gap_stats[n]["avg_gap"],
            "tags": tag
        })
        
    # 2. 純拖牌前 5 名 (含並列)
    sorted_drag = sorted(drag_freq.items(), key=lambda x: x[1], reverse=True)
    cutoff_val = sorted_drag[4][1] if len(sorted_drag) > 4 else 1
    pure_drag_top = [item for item in sorted_drag if item[1] >= cutoff_val and item[1] > 0]
    pure_drag_top.sort(key=lambda x: (x[1], -x[0]), reverse=True)
    
    pure_drag_list = []
    current_rank = 1
    for idx, (num, count) in enumerate(pure_drag_top):
        if idx > 0 and count < pure_drag_top[idx - 1][1]:
            current_rank = idx + 1
        pure_drag_list.append({
            "rank": current_rank,
            "num": f"{num:02d}",
            "count": count
        })
        
    # 3. 歷史拖牌匹配明細 (近 10 次)
    match_list = []
    for m in reversed(matches[-10:]):
        match_list.append({
            "period": m["curr"]["period"],
            "date": m["curr"]["date"],
            "nums": [f"{n:02d}" for n in m["curr"]["nums"]],
            "next_nums": [f"{n:02d}" for n in m["next"]["nums"]],
            "intersect": [f"{n:02d}" for n in m["intersect"]]
        })
        
    # 4. 1-39 完整號碼表
    full_table = []
    for n in range(1, 40):
        full_table.append({
            "num": f"{n:02d}",
            "drag_count": drag_freq[n],
            "count_30": hot_cold_counts[n],
            "current_gap": gap_stats[n]["current_gap"],
            "avg_gap": gap_stats[n]["avg_gap"],
            "max_gap": gap_stats[n]["max_gap"]
        })
        
    # 5. 回測數據
    backtest_stats = {}
    if backtest_data:
        hits_A = [r["hits"] for r in backtest_data["strategy_A"]]
        hits_B = [r["hits"] for r in backtest_data["strategy_B"]]
        
        def get_rates(hits):
            total = len(hits)
            h0 = hits.count(0)
            h1 = hits.count(1)
            h2 = hits.count(2)
            h3 = hits.count(3)
            h4_plus = sum(1 for h in hits if h >= 4)
            return {
                "avg": round(sum(hits)/total, 2),
                "rates": [
                    {"stars": 0, "pct": round(h0/total*100, 1), "count": h0},
                    {"stars": 1, "pct": round(h1/total*100, 1), "count": h1},
                    {"stars": 2, "pct": round(h2/total*100, 1), "count": h2},
                    {"stars": 3, "pct": round(h3/total*100, 1), "count": h3},
                    {"stars": 4, "pct": round(h4_plus/total*100, 1), "count": h4_plus}
                ]
            }
        backtest_stats = {
            "strategy_A": get_rates(hits_A),
            "strategy_B": get_rates(hits_B)
        }
        
    web_data = {
        "update_time": datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "latest_draw": {
            "period": latest_draw["period"],
            "date": latest_draw["date"],
            "nums": [f"{n:02d}" for n in selected_nums]
        },
        "threshold": threshold,
        "hot_nums": [f"{n:02d}" for n in hot_nums],
        "cold_nums": [f"{n:02d}" for n in cold_nums],
        "recommendations": recommendations,
        "pure_drag": pure_drag_list,
        "matches": match_list,
        "full_table": full_table,
        "backtest": backtest_stats
    }
    
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    webapp_dir = os.path.join(base_dir, "webapp")
    if not os.path.exists(webapp_dir):
        os.makedirs(webapp_dir)
        
    json_path = os.path.join(webapp_dir, "data.json")
    try:
        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump(web_data, f, ensure_ascii=False, indent=2)
        print(f"成功生成 Web JSON 數據：{json_path}")
    except Exception as e:
        print(f"錯誤: 無法生成 Web JSON 數據: {e}")

def generate_report(draws, threshold=2):
    """
    執行完整分析流程並生成今彩539分析報告
    """
    if not draws:
        print("沒有開獎資料，無法生成報告。")
        return
        
    latest_draw = draws[-1]
    selected_nums = latest_draw["nums"]
    
    matches, drag_freq = run_drag_analysis(draws, selected_nums, threshold)
    hot_cold_counts, hot_nums, cold_nums = run_hot_cold_analysis(draws, window=30)
    gap_stats = run_gap_analysis(draws)
    backtest_data = backtest_strategies(draws, test_window=30, threshold=threshold)
    
    max_drag = max(drag_freq.values()) if drag_freq.values() else 1
    scores = {}
    for n in range(1, 40):
        s_drag = drag_freq[n] / max_drag
        current_gap = gap_stats[n]["current_gap"]
        avg_gap = gap_stats[n]["avg_gap"] if gap_stats[n]["avg_gap"] > 0 else 7.8
        s_gap = min(current_gap / avg_gap, 2.0) / 2.0
        scores[n] = 0.6 * s_drag + 0.4 * s_gap
        
    top_recommended = sorted(range(1, 40), key=lambda x: scores[x], reverse=True)[:10]
    
    # 產生3碼日期前綴 (基於最新開獎日期)
    try:
        sep = '/' if '/' in latest_draw["date"] else '-'
        parts = latest_draw["date"].split(sep)
        m = int(parts[1])
        d = int(parts[2])
        file_prefix = f"{m}{d:02d}"
    except Exception:
        now = datetime.datetime.now()
        file_prefix = f"{now.month}{now.day:02d}"
    
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    report_path = os.path.join(base_dir, f"{file_prefix}539分析.md")
    
    with open(report_path, mode='w', encoding='utf-8') as f:
        f.write(f"# 今彩539數據分析與預測報告\n\n")
        f.write(f"> **報告生成時間**: {latest_draw['date']} (資料庫最新期別: {latest_draw['period']})\n")
        f.write(f"> **本期分析輸入 (最新獎號)**: " + " ".join(f"{n:02d}" for n in selected_nums) + "\n")
        f.write(f"> **匹配拖牌門檻**: $\\ge$ {threshold} 星\n\n")
        
        f.write(f"## 1. 預測推薦號碼 (綜合評分 Top 10)\n\n")
        f.write(f"結合了**拖牌出現頻率 (權重 60%)** 與 **遺漏值超限程度 (權重 40%)** 的綜合評估推薦：\n\n")
        f.write(f"| 推薦排名 | 號碼 | 綜合分數 | 拖牌次數 | 目前遺漏期數 | 平均間隔期數 | 狀態標籤 |\n")
        f.write(f"| :---: | :---: | :---: | :---: | :---: | :---: | :---: |\n")
        
        for rank, n in enumerate(top_recommended, 1):
            tag = []
            if n in hot_nums:
                tag.append("🔥近期熱門")
            elif n in cold_nums:
                tag.append("❄️近期冷門")
                
            if gap_stats[n]["current_gap"] > gap_stats[n]["avg_gap"] * 1.5:
                tag.append("⚠️遺漏值偏高")
            elif gap_stats[n]["current_gap"] == 0:
                tag.append("🔄上期剛開")
                
            tag_str = "/".join(tag) if tag else "正常"
            f.write(f"| {rank} | **{n:02d}** | {scores[n]:.2f} | {drag_freq[n]} 次 | {gap_stats[n]['current_gap']} 期 | {gap_stats[n]['avg_gap']} 期 | {tag_str} |\n")
            
        f.write(f"\n**純拖牌分析前 5 名 (含第 5 名同次數號碼)**：\n\n")
        sorted_drag = sorted(drag_freq.items(), key=lambda x: x[1], reverse=True)
        cutoff_val = sorted_drag[4][1] if len(sorted_drag) > 4 else 1
        pure_drag_top = [item for item in sorted_drag if item[1] >= cutoff_val and item[1] > 0]
        pure_drag_top.sort(key=lambda x: (x[1], -x[0]), reverse=True)
        current_rank = 1
        for idx, (num, count) in enumerate(pure_drag_top):
            if idx > 0 and count < pure_drag_top[idx - 1][1]:
                current_rank = idx + 1
            f.write(f"* **第 {current_rank} 名**：號碼 **{num:02d}** (歷史拖牌出現 {count} 次)\n")
            
        f.write(f"\n---\n\n## 2. 歷史拖牌匹配明細 (近 10 次匹配)\n\n")
        f.write(f"尋找歷史上開出與本期相同的獎號中，包含 $\\ge$ {threshold} 個號碼的期數，其**下一期**的開獎結果：\n\n")
        f.write(f"| 歷史期別 | 開獎日期 | 當期開獎號碼 (匹配號碼以粗體表示) | 其下一期開獎號碼 |\n")
        f.write(f"| :---: | :---: | :--- | :--- |\n")
        
        for m in reversed(matches[-10:]):
            curr_nums_str = []
            for n in m["curr"]["nums"]:
                if n in selected_nums:
                    curr_nums_str.append(f"**{n:02d}**")
                else:
                    curr_nums_str.append(f"{n:02d}")
            curr_nums_display = " ".join(curr_nums_str)
            next_nums_display = " ".join(f"{n:02d}" for n in m["next"]["nums"])
            f.write(f"| {m['curr']['period']} | {m['curr']['date']} | {curr_nums_display} | {next_nums_display} |\n")
            
        f.write(f"\n*歷史匹配總次數: {len(matches)} 次*\n\n---\n\n")
        
        f.write(f"## 3. 冷熱門與遺漏值分佈 (1-39 完整表)\n\n")
        f.write(f"* **熱門號碼** (近30期開出次數最高前8名): " + ", ".join(f"**{n:02d}**" for n in hot_nums) + "\n")
        f.write(f"* **冷門號碼** (近30期開出次數最低後8名): " + ", ".join(f"**{n:02d}**" for n in cold_nums) + "\n\n")
        
        f.write(f"| 號碼 | 拖牌次數 | 近30期次數 | 目前遺漏期數 | 歷史平均間隔 | 歷史最大間隔 |\n")
        f.write(f"| :---: | :---: | :---: | :---: | :---: | :---: |\n")
        for n in range(1, 40):
            f.write(f"| {n:02d} | {drag_freq[n]} | {hot_cold_counts[n]} | {gap_stats[n]['current_gap']} | {gap_stats[n]['avg_gap']} | {gap_stats[n]['max_gap']} |\n")
            
        f.write(f"\n---\n\n## 4. 預測策略歷史回測分析 (近 30 期表現)\n\n")
        
        if backtest_data:
            hits_A = [r["hits"] for r in backtest_data["strategy_A"]]
            hits_B = [r["hits"] for r in backtest_data["strategy_B"]]
            
            def get_hit_stats(hits):
                total = len(hits)
                h0 = hits.count(0)
                h1 = hits.count(1)
                h2 = hits.count(2)
                h3 = hits.count(3)
                h4_plus = sum(1 for h in hits if h >= 4)
                return {
                    "avg": round(sum(hits)/total, 2),
                    "rates": [
                        f"中 0 星: {h0/total*100:.1f}% ({h0}次)",
                        f"中 1 星: {h1/total*100:.1f}% ({h1}次)",
                        f"中 2 星(普獎): {h2/total*100:.1f}% ({h2}次)",
                        f"中 3 星: {h3/total*100:.1f}% ({h3}次)",
                        f"中 4 星以上: {h4_plus/total*100:.1f}% ({h4_plus}次)"
                    ]
                }
                
            stats_A = get_hit_stats(hits_A)
            stats_B = get_hit_stats(hits_B)
            
            f.write(f"### 預測策略評比：\n\n")
            f.write(f"* **策略 A (純拖牌頻率 Top 5)**:\n")
            f.write(f"  * 歷史平均每期命中: **{stats_A['avg']}** 個號碼\n")
            for rate in stats_A["rates"]:
                f.write(f"  * {rate}\n")
            f.write(f"\n")
            f.write(f"* **策略 B (綜合評分 Top 5 - 拖牌 60% + 遺漏值 40%)**:\n")
            f.write(f"  * 歷史平均每期命中: **{stats_B['avg']}** 個號碼\n")
            for rate in stats_B["rates"]:
                f.write(f"  * {rate}\n")
        else:
            f.write(f"數據庫內期數不足，無法進行 30 期回測。\n")
            
    print(f"成功生成分析報告：{report_path}")
    generate_json(latest_draw, selected_nums, threshold, top_recommended, scores, drag_freq, gap_stats, hot_nums, cold_nums, matches, hot_cold_counts, backtest_data)
    generate_docx(latest_draw, selected_nums, threshold, top_recommended, scores, drag_freq, gap_stats, hot_nums, cold_nums, matches, hot_cold_counts, backtest_data, file_prefix)

if __name__ == "__main__":
    if len(sys.argv) > 1 and sys.argv[1] == "--test":
        print("正在測試載入今彩539開獎數據...")
        draws = load_data()
        print(f"成功載入 {len(draws)} 期開獎數據。")
        if draws:
            print(f"第一期: 期別 {draws[0]['period']}, 日期 {draws[0]['date']}, 獎號 {draws[0]['nums']}")
            print(f"最新一期: 期別 {draws[-1]['period']}, 日期 {draws[-1]['date']}, 獎號 {draws[-1]['nums']}")
            print("測試通過！")
        sys.exit(0)
        
    print("開始執行今彩539數據分析...")
    draws = load_data()
    generate_report(draws, threshold=2)
